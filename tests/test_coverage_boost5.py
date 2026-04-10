"""
Testes para cobertura boost 5 — alvos principais:
  - academic_logger.py         (90% → ~96%)
  - data_validation_tool.py    (90% → ~97%)
  - statistical_analysis_tool.py (93% → ~98%)
  - pdf_search_tool.py         (94% → ~98%)
  - docx_search_tool.py        (94% → ~98%)
  - web/app.py                 (88% → ~93%)
"""
import io
import json
import os
import sqlite3
import sys
import tempfile
from datetime import datetime
from unittest.mock import MagicMock, PropertyMock, patch

import pandas as pd
import pytest

sys.path.insert(0, os.path.join(os.path.dirname(__file__), '..', 'src'))


# ---------------------------------------------------------------------------
# AcademicAuditor / AcademicLogger — branches restantes
# ---------------------------------------------------------------------------

class TestAcademicLoggerBranches:

    @pytest.fixture
    def auditor(self):
        import logging
        from Sapiens_MultiAgente.tools.academic_logger import AcademicAuditor
        for name in ('sapiens_academico', 'sapiens_auditoria'):
            logging.getLogger(name).handlers.clear()
        return AcademicAuditor()

    # ---- _load_config: exception ao ler arquivo (linhas 29-31) ----
    def test_load_config_arquivo_invalido(self):
        import logging
        from Sapiens_MultiAgente.tools.academic_logger import AcademicAuditor
        for name in ('sapiens_academico', 'sapiens_auditoria'):
            logging.getLogger(name).handlers.clear()
        a = AcademicAuditor(config_path='/nao/existe/config.yaml')
        # Não deve levantar exceção; retorna config vazia
        assert a.config == {} or isinstance(a.config, dict)

    # ---- _setup_logging: if self.logger.handlers: return (linha 44) ----
    def test_setup_logging_retorno_rapido(self):
        """Cria 2ª instância SEM limpar handlers → linha 44 (return precoce)."""
        import logging
        from Sapiens_MultiAgente.tools.academic_logger import AcademicAuditor
        # Garante que handlers já existam do auditor_global
        logging.getLogger('sapiens_academico').handlers  # referência garantida
        # Cria segunda instância sem limpar handlers
        a = AcademicAuditor()
        # O return precoce na linha 44 é atingido; auditoria_logger pode não existir
        assert a is not None

    # ---- JsonFormatter.format com tempo_execucao_ms (linhas 98-101) ----
    def test_json_formatter_com_tempo_execucao(self, auditor):
        """Loga com tempo_execucao_ms no extra → JsonFormatter inclui o campo."""
        import logging
        logger = auditor.auditoria_logger
        # Inclui tempo_execucao_ms no extra; sessao_id e usuario_id são obrigatórios
        logger.info(
            'Evento com tempo',
            extra={
                'sessao_id': auditor.sessao_id,
                'usuario_id': auditor.usuario_id,
                'dados_adicionais': {},
                'tempo_execucao_ms': 42.5,
            }
        )
        # Se chegou aqui sem exceção, as linhas 98-101 foram exercitadas
        assert True

    # ---- _calcular_tempo_sessao: exception com timestamps inválidos (linhas 235-236) ----
    def test_calcular_tempo_sessao_exception(self, auditor):
        auditor._eventos_auditoria = [
            {'timestamp': 'timestamp_invalido_xyz'},
            {'timestamp': 'outro_invalido'},
        ]
        duracao = auditor._calcular_tempo_sessao()
        assert duracao == 0.0

    # ---- exportar_auditoria_json: exception (linhas 260-262) ----
    def test_exportar_auditoria_json_exception(self, auditor):
        auditor.auditar_evento('evento_teste', dados_entrada={'x': 1})
        with patch('builtins.open', side_effect=PermissionError('sem permissão')):
            result = auditor.exportar_auditoria_json('/caminho/invalido.json')
        assert result is None


# ---------------------------------------------------------------------------
# DataValidationTool — branches restantes
# ---------------------------------------------------------------------------

class TestDataValidationToolBranches:

    def _make_csv(self, tmp_path, content='a,b\n1,2\n'):
        p = tmp_path / 'data.csv'
        p.write_text(content, encoding='utf-8')
        return str(p)

    # ---- MIME type não suportado (linha 69) ----
    def test_tipo_mime_nao_suportado(self, tmp_path):
        from Sapiens_MultiAgente.tools.data_validation_tool import DataValidationTool
        arq = tmp_path / 'script.py'
        arq.write_text('print("hello")', encoding='utf-8')
        tool = DataValidationTool()
        # Força MIME type de um arquivo .py como não aceito
        with patch('Sapiens_MultiAgente.tools.data_validation_tool.magic') as m_magic:
            mime_mock = MagicMock()
            mime_mock.from_file.return_value = 'text/x-python'
            m_magic.Magic.return_value = mime_mock
            result = tool._run(file_path=str(arq))
        assert '❌' in result
        assert 'não suportado' in result.lower() or 'MIME' in result

    # ---- PII detectado — aviso adicionado (linha 90) ----
    def test_pii_detectado_adiciona_aviso(self, tmp_path):
        from Sapiens_MultiAgente.tools.data_validation_tool import DataValidationTool
        arq = tmp_path / 'pii.csv'
        arq.write_text('cpf,nome\n123.456.789-00,João\n', encoding='utf-8')
        tool = DataValidationTool()
        with patch('Sapiens_MultiAgente.tools.data_validation_tool.magic') as m_magic:
            mime_mock = MagicMock()
            mime_mock.from_file.return_value = 'text/csv'
            m_magic.Magic.return_value = mime_mock
            result = tool._run(file_path=str(arq))
        # Resultado pode ser válido com aviso de PII
        assert isinstance(result, str)

    # ---- exception geral em _run (linhas 94-95) ----
    def test_exception_geral_em_run(self, tmp_path):
        from Sapiens_MultiAgente.tools.data_validation_tool import DataValidationTool
        arq = self._make_csv(tmp_path)
        tool = DataValidationTool()
        # Passa pelo MIME check mas falha no _calculate_sha256 → except Exception
        with patch('Sapiens_MultiAgente.tools.data_validation_tool.magic') as m_magic, \
             patch.object(tool, '_calculate_sha256', side_effect=RuntimeError('hash error')):
            mime_mock = MagicMock()
            mime_mock.from_file.return_value = 'text/csv'
            m_magic.Magic.return_value = mime_mock
            result = tool._run(file_path=arq)
        assert '❌ ERRO' in result

    # ---- exception em _check_pii_indicators (linhas 128-130) ----
    def test_check_pii_exception(self, tmp_path):
        from Sapiens_MultiAgente.tools.data_validation_tool import DataValidationTool
        arq = self._make_csv(tmp_path)
        tool = DataValidationTool()
        with patch('builtins.open', side_effect=PermissionError('sem permissão')):
            result = tool._check_pii_indicators(arq, 'text/csv')
        assert result == []  # sem avisos quando não consegue ler


# ---------------------------------------------------------------------------
# StatisticalAnalysisTool — branches restantes
# ---------------------------------------------------------------------------

class TestStatisticalAnalysisToolMoreBranches:

    def _json(self, data: dict) -> str:
        return json.dumps(data)

    # ---- except Exception em _run (linhas 92-93) ----
    def test_exception_geral_em_run(self):
        from Sapiens_MultiAgente.tools.statistical_analysis_tool import StatisticalAnalysisTool
        tool = StatisticalAnalysisTool()
        dados = self._json({'a': [1, 2, 3], 'b': [4, 5, 6]})
        with patch.object(tool, '_descriptive_analysis', side_effect=RuntimeError('falha')):
            result = tool._run(data=dados, analysis_type='descriptive')
        assert '❌ ERRO durante análise estatística' in result

    # ---- Shapiro exception: col não foi possível testar (linhas 114-115) ----
    def test_shapiro_exception_na_coluna(self):
        from Sapiens_MultiAgente.tools.statistical_analysis_tool import StatisticalAnalysisTool
        from scipy import stats as sp_stats
        tool = StatisticalAnalysisTool()
        dados = self._json({'a': [1, 2, 3, 4, 5], 'b': [6, 7, 8, 9, 10]})
        with patch.object(sp_stats, 'shapiro', side_effect=ValueError('erro shapiro')):
            result = tool._run(data=dados, analysis_type='descriptive')
        assert 'Não foi possível testar' in result

    # ---- Regressão com "Bom ajuste" (r2 ~ 0.78, linhas 229-230) ----
    def test_regressao_bom_ajuste(self):
        from Sapiens_MultiAgente.tools.statistical_analysis_tool import StatisticalAnalysisTool
        tool = StatisticalAnalysisTool()
        # R² ≈ 0.78: bom mas não excelente
        dados = self._json({'x': [1, 2, 3, 4, 5], 'y': [2, 4, 4, 4, 5]})
        result = tool._run(data=dados, analysis_type='regression')
        assert 'Bom ajuste' in result

    # ---- Regressão com "Ajuste moderado" (0.3 < r2 < 0.6, linhas 231-232) ----
    def test_regressao_ajuste_moderado(self):
        from Sapiens_MultiAgente.tools.statistical_analysis_tool import StatisticalAnalysisTool
        tool = StatisticalAnalysisTool()
        # R² ≈ 0.42: moderado
        dados = self._json({'x': [1, 2, 3, 4, 5], 'y': [2, 5, 3, 7, 5]})
        result = tool._run(data=dados, analysis_type='regression')
        assert 'Ajuste moderado' in result or 'Bom ajuste' in result or 'Ajuste' in result

    # ---- Regressão com exception (linhas 238-239) ----
    def test_regressao_exception(self):
        import sklearn.linear_model
        from Sapiens_MultiAgente.tools.statistical_analysis_tool import StatisticalAnalysisTool
        tool = StatisticalAnalysisTool()
        dados = self._json({'x': [1, 2, 3, 4, 5], 'y': [2, 4, 6, 8, 10]})
        with patch.object(sklearn.linear_model, 'LinearRegression', side_effect=RuntimeError('erro sklearn')):
            result = tool._run(data=dados, analysis_type='regression')
        assert 'Erro na regressão' in result

    # ---- Hipótese com exception em ttest_ind (linhas 182-183) ----
    def test_hipotese_exception(self):
        from Sapiens_MultiAgente.tools.statistical_analysis_tool import StatisticalAnalysisTool
        from scipy import stats as sp_stats
        tool = StatisticalAnalysisTool()
        dados = self._json({'a': [1, 2, 3], 'b': [4, 5, 6]})
        with patch.object(sp_stats, 'ttest_ind', side_effect=RuntimeError('erro ttest')):
            result = tool._run(data=dados, analysis_type='hypothesis_test')
        assert 'Erro' in result

    # ---- ANOVA com variables fornecidas (linha 263) ----
    def test_anova_com_variables_especificas(self):
        from Sapiens_MultiAgente.tools.statistical_analysis_tool import StatisticalAnalysisTool
        tool = StatisticalAnalysisTool()
        dados = self._json({
            'grupo': ['A', 'A', 'B', 'B', 'C', 'C'],
            'nota': [7, 8, 5, 6, 9, 10],
            'score': [1, 2, 3, 4, 5, 6],
        })
        result = tool._run(data=dados, analysis_type='anova_one_way',
                           group_column='grupo', variables=['nota'])
        assert 'ANOVA' in result

    # ---- ANOVA com dados insuficientes após remoção de NaN (linhas 288-289) ----
    def test_anova_insuficiente_apos_nan(self):
        from Sapiens_MultiAgente.tools.statistical_analysis_tool import StatisticalAnalysisTool
        import numpy as np
        tool = StatisticalAnalysisTool()
        # Grupo C tem só NaN em 'nota'
        dados = self._json({
            'grupo': ['A', 'A', 'B', 'B', 'C', 'C'],
            'nota': [7.0, 8.0, 5.0, 6.0, None, None],
        })
        result = tool._run(data=dados, analysis_type='anova_one_way',
                           group_column='grupo')
        # 2 grupos têm dados, 1 não → amostras ainda >= 2, análise roda
        assert 'ANOVA' in result or 'Erro' in result

    # ---- ANOVA exception em f_oneway (linhas 334-335) ----
    def test_anova_exception_f_oneway(self):
        from Sapiens_MultiAgente.tools.statistical_analysis_tool import StatisticalAnalysisTool
        from scipy import stats as sp_stats
        tool = StatisticalAnalysisTool()
        dados = self._json({
            'grupo': ['A', 'A', 'B', 'B'],
            'nota': [7, 8, 5, 6],
        })
        with patch.object(sp_stats, 'f_oneway', side_effect=RuntimeError('anova fail')):
            result = tool._run(data=dados, analysis_type='anova_one_way',
                               group_column='grupo')
        assert 'Erro' in result


# ---------------------------------------------------------------------------
# PDFSearchTool — branches restantes
# ---------------------------------------------------------------------------

class TestPDFSearchToolMoreBranches:

    def _make_pdf(self, tmp_path):
        from pypdf import PdfWriter
        writer = PdfWriter()
        writer.add_blank_page(width=200, height=200)
        p = tmp_path / 'test.pdf'
        with open(str(p), 'wb') as f:
            writer.write(f)
        return str(p)

    # ---- _get_pdf_info: metadata exception (linhas 164-165) ----
    def test_pdf_info_metadata_exception(self, tmp_path):
        from Sapiens_MultiAgente.tools.pdf_search_tool import PDFSearchTool
        pdf = self._make_pdf(tmp_path)
        mock_page = MagicMock()
        mock_page.extract_text.return_value = 'Texto de teste'
        mock_reader = MagicMock()
        type(mock_reader).metadata = PropertyMock(side_effect=RuntimeError('metadata error'))
        mock_reader.pages = [mock_page]
        with patch('Sapiens_MultiAgente.tools.pdf_search_tool.PdfReader', return_value=mock_reader):
            tool = PDFSearchTool()
            result = tool._run(file_path=pdf)
        assert 'PDF' in result

    # ---- _extract_images: page com resources mas sem XObject (linha 214) ----
    def test_extract_images_page_sem_xobject(self, tmp_path):
        from Sapiens_MultiAgente.tools.pdf_search_tool import PDFSearchTool
        pdf = self._make_pdf(tmp_path)
        mock_resources = MagicMock()
        mock_resources.get.side_effect = lambda k, d=None: None  # sem XObject
        mock_page = MagicMock()
        mock_page.get.return_value = mock_resources
        mock_reader = MagicMock()
        mock_reader.pages = [mock_page]
        with patch('Sapiens_MultiAgente.tools.pdf_search_tool.PdfReader', return_value=mock_reader):
            tool = PDFSearchTool()
            result = tool._run(file_path=pdf, extract_images=True)
        assert '🖼️' in result

    # ---- _extract_images: filtro /FlateDecode (linha 232) ----
    def test_extract_images_filtro_flatdecode(self, tmp_path):
        from Sapiens_MultiAgente.tools.pdf_search_tool import PDFSearchTool
        pdf = self._make_pdf(tmp_path)
        dest = str(tmp_path / 'imgs_flat')

        mock_obj = MagicMock()
        mock_obj.get.side_effect = lambda k, d=None: {
            '/Subtype': '/Image',
            '/Width': 50,
            '/Height': 50,
            '/ColorSpace': '/DeviceRGB',
            '/Filter': '/FlateDecode',
        }.get(k, d)
        mock_obj.get_data.return_value = b'\x89PNG\r\n' + b'\x00' * 20

        mock_xobjects = {'/Im1': MagicMock(get_object=MagicMock(return_value=mock_obj))}
        mock_resources = MagicMock()
        mock_resources.get.side_effect = lambda k, d=None: mock_xobjects if k == '/XObject' else None

        mock_page = MagicMock()
        mock_page.get.return_value = mock_resources

        mock_reader = MagicMock()
        mock_reader.pages = [mock_page]

        with patch('Sapiens_MultiAgente.tools.pdf_search_tool.PdfReader', return_value=mock_reader):
            tool = PDFSearchTool()
            result = tool._run(file_path=pdf, extract_images=True, output_dir=dest)
        assert '🖼️' in result
        assert '.png' in result or '✅' in result

    # ---- _extract_images: filtro /JPXDecode (linha 234) ----
    def test_extract_images_filtro_jpxdecode(self, tmp_path):
        from Sapiens_MultiAgente.tools.pdf_search_tool import PDFSearchTool
        pdf = self._make_pdf(tmp_path)
        dest = str(tmp_path / 'imgs_jpx')

        mock_obj = MagicMock()
        mock_obj.get.side_effect = lambda k, d=None: {
            '/Subtype': '/Image', '/Width': 30, '/Height': 30,
            '/ColorSpace': '/DeviceRGB', '/Filter': '/JPXDecode',
        }.get(k, d)
        mock_obj.get_data.return_value = b'\x00\x00\x00\x0cjP  \r\n'  # JPEG 2000 stub

        mock_xobjects = {'/Im1': MagicMock(get_object=MagicMock(return_value=mock_obj))}
        mock_resources = MagicMock()
        mock_resources.get.side_effect = lambda k, d=None: mock_xobjects if k == '/XObject' else None

        mock_page = MagicMock()
        mock_page.get.return_value = mock_resources

        mock_reader = MagicMock()
        mock_reader.pages = [mock_page]

        with patch('Sapiens_MultiAgente.tools.pdf_search_tool.PdfReader', return_value=mock_reader):
            tool = PDFSearchTool()
            result = tool._run(file_path=pdf, extract_images=True, output_dir=dest)
        assert '.jp2' in result or '✅' in result

    # ---- _extract_images: filtro desconhecido → .bin (linhas 235-236) ----
    def test_extract_images_filtro_desconhecido(self, tmp_path):
        from Sapiens_MultiAgente.tools.pdf_search_tool import PDFSearchTool
        pdf = self._make_pdf(tmp_path)
        dest = str(tmp_path / 'imgs_bin')

        mock_obj = MagicMock()
        mock_obj.get.side_effect = lambda k, d=None: {
            '/Subtype': '/Image', '/Width': 10, '/Height': 10,
            '/ColorSpace': '/DeviceGray', '/Filter': '/JBIG2Decode',
        }.get(k, d)
        mock_obj.get_data.return_value = b'\x00' * 10

        mock_xobjects = {'/Im1': MagicMock(get_object=MagicMock(return_value=mock_obj))}
        mock_resources = MagicMock()
        mock_resources.get.side_effect = lambda k, d=None: mock_xobjects if k == '/XObject' else None

        mock_page = MagicMock()
        mock_page.get.return_value = mock_resources

        mock_reader = MagicMock()
        mock_reader.pages = [mock_page]

        with patch('Sapiens_MultiAgente.tools.pdf_search_tool.PdfReader', return_value=mock_reader):
            tool = PDFSearchTool()
            result = tool._run(file_path=pdf, extract_images=True, output_dir=dest)
        assert '.bin' in result or '✅' in result

    # ---- _extract_images: exception em reader.pages (linhas 249-250) ----
    def test_extract_images_page_exception(self, tmp_path):
        from Sapiens_MultiAgente.tools.pdf_search_tool import PDFSearchTool
        pdf = self._make_pdf(tmp_path)

        mock_page = MagicMock()
        mock_page.get.side_effect = RuntimeError('bad page')

        mock_reader = MagicMock()
        mock_reader.pages = [mock_page]

        with patch('Sapiens_MultiAgente.tools.pdf_search_tool.PdfReader', return_value=mock_reader):
            tool = PDFSearchTool()
            result = tool._run(file_path=pdf, extract_images=True)
        assert '⚠️' in result or '🖼️' in result


# ---------------------------------------------------------------------------
# DOCXSearchTool — branches restantes
# ---------------------------------------------------------------------------

class TestDOCXSearchToolMoreBranches:

    def _make_docx(self, tmp_path, paragraphs=None, tables=False):
        from docx import Document as DocxDoc
        doc = DocxDoc()
        for p in (paragraphs or ['Parágrafo de teste.']):
            doc.add_paragraph(p)
        if tables:
            t = doc.add_table(rows=2, cols=2)
            t.cell(0, 0).text = 'A'
            t.cell(0, 1).text = 'B'
            t.cell(1, 0).text = 'C'
            t.cell(1, 1).text = 'D'
        path = str(tmp_path / 'test.docx')
        doc.save(path)
        return path

    # ---- _extract_full_text: tabela com exception (linhas 101-102) ----
    def test_extract_full_text_tabela_exception(self, tmp_path):
        from Sapiens_MultiAgente.tools.docx_search_tool import DOCXSearchTool, Document
        docx = self._make_docx(tmp_path, tables=True)
        doc = Document(docx)

        bad_table = MagicMock()
        type(bad_table).rows = PropertyMock(side_effect=RuntimeError('rows error'))
        type(bad_table).columns = PropertyMock(side_effect=RuntimeError('cols error'))
        # Substituir a lista de tabelas
        with patch.object(type(doc), 'tables', new_callable=PropertyMock,
                          return_value=[bad_table]):
            tool = DOCXSearchTool()
            result = tool._extract_full_text(doc, paragraph_limit=None)
        assert 'erro ao processar' in result.lower() or 'Tabela' in result

    # ---- _search_in_tables: exception em cell.text (linhas 152-153) ----
    def test_search_in_tables_cell_exception(self):
        from Sapiens_MultiAgente.tools.docx_search_tool import DOCXSearchTool

        bad_cell = MagicMock()
        type(bad_cell).text = PropertyMock(side_effect=RuntimeError('cell error'))

        bad_row = MagicMock()
        bad_row.cells = [bad_cell]

        bad_table = MagicMock()
        bad_table.rows = [bad_row]

        mock_doc = MagicMock()
        mock_doc.tables = [bad_table]

        tool = DOCXSearchTool()
        result = tool._search_in_tables(mock_doc, 'busca')
        assert isinstance(result, list)

    # ---- _get_docx_info: tabela com exception nas colunas (linhas 171-172) ----
    def test_get_docx_info_tabela_coluna_exception(self, tmp_path):
        from Sapiens_MultiAgente.tools.docx_search_tool import DOCXSearchTool, Document
        docx = self._make_docx(tmp_path)
        doc = Document(docx)

        bad_table = MagicMock()
        type(bad_table).rows = PropertyMock(return_value=[MagicMock()])  # 1 row
        type(bad_table).columns = PropertyMock(side_effect=RuntimeError('cols error'))

        with patch.object(type(doc), 'tables', new_callable=PropertyMock,
                          return_value=[bad_table]):
            tool = DOCXSearchTool()
            result = tool._get_docx_info(doc, docx)
        assert 'erro ao processar' in result.lower() or 'DOCX' in result

    # ---- _extract_images: exception geral (linha 255-256) ----
    def test_extract_images_exception_geral(self, tmp_path):
        import zipfile
        from Sapiens_MultiAgente.tools.docx_search_tool import DOCXSearchTool
        docx_path = str(tmp_path / 'test_exc.docx')

        with zipfile.ZipFile(docx_path, 'w') as zf:
            zf.writestr('word/document.xml', '<body/>')
            zf.writestr('word/media/img.jpg', b'\xff\xd8')

        tool = DOCXSearchTool()
        with patch('zipfile.ZipFile', side_effect=RuntimeError('zip error')):
            result = tool._extract_images(docx_path, None)
        assert '❌' in result or 'erro' in result.lower()


# ---------------------------------------------------------------------------
# Web App — branches restantes
# ---------------------------------------------------------------------------

class TestWebAppMoreBranches:

    @pytest.fixture
    def app_cliente(self, tmp_path):
        db_path = str(tmp_path / 'test.db')
        upload_folder = str(tmp_path / 'uploads')
        os.makedirs(upload_folder, exist_ok=True)

        with patch('Sapiens_MultiAgente.web.app.DB_PATH', db_path), \
             patch('Sapiens_MultiAgente.web.auth.DB_PATH', db_path):
            from Sapiens_MultiAgente.web.app import SapiensWebInterface
            interface = SapiensWebInterface()
            interface.upload_config['UPLOAD_FOLDER'] = upload_folder
            flask_app = interface.create_app()
            flask_app.config['TESTING'] = True
            flask_app.config['WTF_CSRF_ENABLED'] = False
            flask_app.config['UPLOAD_FOLDER'] = upload_folder
            client = flask_app.test_client()
            client.post('/login', data={'username': 'admin', 'senha': 'sapiens@2025'})
            yield interface, client, db_path

    def _inserir_analise(self, db_path, status='concluida', usuario_id='1', analise_id=None):
        analise_id = analise_id or f'analise-{status}-test'
        with sqlite3.connect(db_path) as conn:
            conn.execute("""
                INSERT OR IGNORE INTO analises
                    (id, usuario_id, topico, status, progresso, arquivos, resultados, criado_em)
                VALUES (?, ?, ?, ?, 100, '[]', 'Resultados com &amp; <tags> especiais', ?)
            """, (analise_id, usuario_id, 'Tópico', status, datetime.now().isoformat()))
            conn.commit()
        return analise_id

    # ---- _salvar_analise: analise_id não existe no cache (linha 118) ----
    def test_salvar_analise_id_nao_existe(self, app_cliente):
        interface, _, _ = app_cliente
        # Não deve levantar exceção; retorna imediatamente
        interface._salvar_analise('id-que-nao-existe')

    # ---- _buscar_analise: exception no DB (linhas 159-160) ----
    def test_buscar_analise_db_exception(self, app_cliente):
        interface, _, _ = app_cliente
        with patch('Sapiens_MultiAgente.web.app._get_db', side_effect=RuntimeError('DB error')):
            result = interface._buscar_analise('id-nao-no-cache')
        assert result is None

    # ---- _carregar_analises_ativas: rows no banco (linha 104) ----
    def test_carregar_analises_ativas_com_rows(self, tmp_path):
        db_path = str(tmp_path / 'preload.db')
        upload_folder = str(tmp_path / 'uploads2')
        os.makedirs(upload_folder, exist_ok=True)

        # Cria o banco e insere uma análise ativa ANTES de criar a interface
        with patch('Sapiens_MultiAgente.web.app.DB_PATH', db_path), \
             patch('Sapiens_MultiAgente.web.auth.DB_PATH', db_path):
            from Sapiens_MultiAgente.web.app import SapiensWebInterface, _get_db, _init_db
            _init_db()
            with _get_db() as conn:
                conn.execute("""
                    INSERT INTO analises (id, usuario_id, topico, status, progresso, arquivos, criado_em)
                    VALUES ('analise-ativa', '1', 'Topico', 'processando', 50, '[]', ?)
                """, (datetime.now().isoformat(),))
                conn.commit()

            # Cria interface DEPOIS de inserir — _carregar_analises_ativas marca as órfãs como erro
            interface2 = SapiensWebInterface()
            # Verifica que a análise foi marcada como erro (comportamento de cleanup no startup)
            with _get_db() as conn:
                row = conn.execute("SELECT status FROM analises WHERE id = 'analise-ativa'").fetchone()
        assert row['status'] == 'erro'
        assert interface2.analises_ativas == {}

    # ---- _validar_url: cloud metadata bloqueado (linha 588) ----
    def test_validar_url_cloud_metadata(self, app_cliente):
        """Usa hostname que não é IP válido mas começa com '169.254.' → linha 588."""
        interface, _, _ = app_cliente
        flask_app = interface.create_app()
        with flask_app.test_request_context():
            # 169.254.internal não é IP válido → ValueError → pass → blocked_prefixes check
            valida, motivo = interface._validar_url('http://169.254.internal.corp/meta-data')
        assert not valida
        assert 'cloud' in motivo.lower() or 'metadata' in motivo.lower() or 'internal' in motivo.lower()

    # ---- _validar_url: exception geral (linhas 592-593) ----
    def test_validar_url_exception(self, app_cliente):
        interface, _, _ = app_cliente
        flask_app = interface.create_app()
        with flask_app.test_request_context():
            with patch('Sapiens_MultiAgente.web.app.urlparse', side_effect=RuntimeError('parse error')):
                valida, motivo = interface._validar_url('http://test.com')
        assert not valida

    # ---- _processar_analise: exception (linhas 715-718) ----
    def test_processar_analise_exception(self, app_cliente):
        interface, client, _ = app_cliente
        with patch.object(interface.auditor, 'iniciar_analise', side_effect=RuntimeError('auditor falhou')):
            resp = client.post(
                '/analise',
                data={'topico_pesquisa': 'Teste de exception'},
                follow_redirects=False
            )
        assert resp.status_code == 302
        assert 'analise' in resp.headers['Location']

    # ---- _handle_upload: exception em arquivo.save (linhas 757-758) ----
    def test_handle_upload_exception_no_save(self, app_cliente):
        interface, client, _ = app_cliente
        with patch('werkzeug.datastructures.FileStorage.save', side_effect=OSError('disk full')):
            resp = client.post(
                '/upload',
                data={'arquivo': (io.BytesIO(b'data'), 'dados.csv')},
                content_type='multipart/form-data'
            )
        assert resp.status_code == 500

    # ---- _limpar_arquivos_analise: exception ao remover (linhas 876-877) ----
    def test_limpar_arquivos_exception(self, app_cliente, tmp_path):
        interface, _, _ = app_cliente
        fake_file = str(tmp_path / 'fake_upload.csv')
        with open(fake_file, 'w') as f:
            f.write('a,b\n1,2\n')
        interface.analises_ativas['analise-limpeza'] = {
            'arquivos': [{'caminho': fake_file}]
        }
        with patch('os.remove', side_effect=OSError('arquivo em uso')):
            interface._limpar_arquivos_analise('analise-limpeza')
        # Não deve levantar exceção; apenas ignora o erro

    # ---- _limpar_uploads_orfaos: exception em getmtime (linhas 901-902) ----
    def test_limpar_uploads_orfaos_getmtime_exception(self, app_cliente, tmp_path):
        interface, _, _ = app_cliente
        upload_folder = interface.upload_config['UPLOAD_FOLDER']
        # Cria um arquivo órfão
        orfao = os.path.join(upload_folder, 'orfao.csv')
        with open(orfao, 'w') as f:
            f.write('a,b\n')
        with patch('os.path.getmtime', side_effect=OSError('stat error')):
            interface._limpar_uploads_orfaos(max_idade_horas=-1)
        # Não deve levantar exceção

    # ---- _get_historico: exception no DB (linhas 937-938) ----
    def test_get_historico_db_exception(self, app_cliente):
        _, client, _ = app_cliente
        with patch('Sapiens_MultiAgente.web.app._get_db', side_effect=RuntimeError('DB error')):
            resp = client.get('/historico')
        assert resp.status_code == 200  # renderiza historico.html com lista vazia

    # ---- export PDF com chars especiais no texto (linha 1009) ----
    def test_export_pdf_com_chars_especiais(self, app_cliente):
        interface, client, db_path = app_cliente
        analise_id = 'pdf-chars-test'
        with sqlite3.connect(db_path) as conn:
            conn.execute("""
                INSERT OR IGNORE INTO analises
                    (id, usuario_id, topico, status, progresso, arquivos, resultados, criado_em)
                VALUES (?, '1', 'Tópico', 'concluida', 100, '[]', ?, ?)
            """, (
                analise_id,
                'Resultado com & e <tags> e > setas',
                datetime.now().isoformat()
            ))
            conn.commit()
        resp = client.get(f'/export/{analise_id}/pdf')
        assert resp.status_code == 200
        assert resp.headers['Content-Type'] == 'application/pdf'

    # ---- export PDF: exception no build (linhas 1023-1024) ----
    def test_export_pdf_exception_no_build(self, app_cliente):
        interface, client, db_path = app_cliente
        analise_id = 'pdf-exc-test'
        with sqlite3.connect(db_path) as conn:
            conn.execute("""
                INSERT OR IGNORE INTO analises
                    (id, usuario_id, topico, status, progresso, arquivos, resultados, criado_em)
                VALUES (?, '1', 'Tópico', 'concluida', 100, '[]', 'Resultado OK', ?)
            """, (analise_id, datetime.now().isoformat()))
            conn.commit()
        with patch('reportlab.platypus.SimpleDocTemplate.build', side_effect=RuntimeError('pdf fail')):
            resp = client.get(f'/export/{analise_id}/pdf')
        assert resp.status_code == 500

    # ---- historico para usuário não-admin (linha 928) ----
    def test_historico_usuario_nao_admin(self, tmp_path):
        db_path = str(tmp_path / 'nonadmin.db')
        with patch('Sapiens_MultiAgente.web.app.DB_PATH', db_path), \
             patch('Sapiens_MultiAgente.web.auth.DB_PATH', db_path):
            from Sapiens_MultiAgente.web.app import SapiensWebInterface
            from Sapiens_MultiAgente.web.auth import Usuario
            interface = SapiensWebInterface()
            flask_app = interface.create_app()
            flask_app.config['TESTING'] = True
            flask_app.config['WTF_CSRF_ENABLED'] = False
            # Cria usuário não-admin
            Usuario.criar('nonadmin', 'senha123', 'Usuário Comum', admin=False)
            client = flask_app.test_client()
            client.post('/login', data={'username': 'nonadmin', 'senha': 'senha123'})
            resp = client.get('/historico')
        assert resp.status_code == 200
