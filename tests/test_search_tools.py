"""
Testes para CSVSearchTool, PDFSearchTool e DOCXSearchTool.
Cobre as lacunas de cobertura: busca, filtros, extração de texto.
"""

import os
import sys
import tempfile
import pytest
from unittest.mock import MagicMock, patch

sys.path.insert(0, os.path.join(os.path.dirname(__file__), '..', 'src'))


# ---------------------------------------------------------------------------
# CSVSearchTool (21% → ~80%)
# ---------------------------------------------------------------------------

class TestCSVSearchTool:
    def _make_csv(self, content: str) -> str:
        f = tempfile.NamedTemporaryFile(mode='w', suffix='.csv', delete=False, encoding='utf-8')
        f.write(content)
        f.flush()
        return f.name

    CSV_BASICO = "nome,nota,curso\nAlice,8.5,Computação\nBob,7.0,Matemática\nCarla,9.2,Computação\n"

    def test_arquivo_inexistente(self):
        from Sapiens_MultiAgente.tools.csv_search_tool import CSVSearchTool
        result = CSVSearchTool()._run(file_path="/nao/existe.csv")
        assert "ERRO" in result

    def test_arquivo_vazio(self):
        from Sapiens_MultiAgente.tools.csv_search_tool import CSVSearchTool
        path = self._make_csv("")
        try:
            result = CSVSearchTool()._run(file_path=path)
            assert "ERRO" in result
        finally:
            os.unlink(path)

    def test_busca_sem_filtros_retorna_todos(self):
        from Sapiens_MultiAgente.tools.csv_search_tool import CSVSearchTool
        path = self._make_csv(self.CSV_BASICO)
        try:
            result = CSVSearchTool()._run(file_path=path)
            assert "RESULTADOS DA BUSCA" in result
            assert "Alice" in result
            assert "Bob" in result
        finally:
            os.unlink(path)

    def test_busca_texto_case_insensitive(self):
        from Sapiens_MultiAgente.tools.csv_search_tool import CSVSearchTool
        path = self._make_csv(self.CSV_BASICO)
        try:
            result = CSVSearchTool()._run(file_path=path, search_query="computação")
            assert "Alice" in result
            assert "Carla" in result
        finally:
            os.unlink(path)

    def test_busca_texto_case_sensitive(self):
        from Sapiens_MultiAgente.tools.csv_search_tool import CSVSearchTool
        path = self._make_csv(self.CSV_BASICO)
        try:
            result = CSVSearchTool()._run(file_path=path, search_query="Computação", case_sensitive=True)
            assert "Alice" in result
        finally:
            os.unlink(path)

    def test_busca_texto_sem_resultados(self):
        from Sapiens_MultiAgente.tools.csv_search_tool import CSVSearchTool
        path = self._make_csv(self.CSV_BASICO)
        try:
            result = CSVSearchTool()._run(file_path=path, search_query="XYZXYZ_inexistente")
            assert "NENHUM RESULTADO" in result
        finally:
            os.unlink(path)

    def test_filtro_coluna_case_insensitive(self):
        from Sapiens_MultiAgente.tools.csv_search_tool import CSVSearchTool
        path = self._make_csv(self.CSV_BASICO)
        try:
            result = CSVSearchTool()._run(file_path=path, column_filters={"curso": "computação"})
            assert "Alice" in result
            assert "Carla" in result
        finally:
            os.unlink(path)

    def test_filtro_coluna_case_sensitive(self):
        from Sapiens_MultiAgente.tools.csv_search_tool import CSVSearchTool
        path = self._make_csv(self.CSV_BASICO)
        try:
            result = CSVSearchTool()._run(file_path=path, column_filters={"nome": "Alice"}, case_sensitive=True)
            assert "Alice" in result
        finally:
            os.unlink(path)

    def test_filtro_numerico_maior(self):
        from Sapiens_MultiAgente.tools.csv_search_tool import CSVSearchTool
        path = self._make_csv(self.CSV_BASICO)
        try:
            result = CSVSearchTool()._run(file_path=path, numeric_filters={"nota": ">8.0"})
            assert "Alice" in result
            assert "Carla" in result
        finally:
            os.unlink(path)

    def test_filtro_numerico_menor(self):
        from Sapiens_MultiAgente.tools.csv_search_tool import CSVSearchTool
        path = self._make_csv(self.CSV_BASICO)
        try:
            result = CSVSearchTool()._run(file_path=path, numeric_filters={"nota": "<8.0"})
            assert "Bob" in result
        finally:
            os.unlink(path)

    def test_filtro_numerico_maior_igual(self):
        from Sapiens_MultiAgente.tools.csv_search_tool import CSVSearchTool
        path = self._make_csv(self.CSV_BASICO)
        try:
            result = CSVSearchTool()._run(file_path=path, numeric_filters={"nota": ">=9.2"})
            assert "Carla" in result
        finally:
            os.unlink(path)

    def test_filtro_numerico_menor_igual(self):
        from Sapiens_MultiAgente.tools.csv_search_tool import CSVSearchTool
        path = self._make_csv(self.CSV_BASICO)
        try:
            result = CSVSearchTool()._run(file_path=path, numeric_filters={"nota": "<=7.0"})
            assert "Bob" in result
        finally:
            os.unlink(path)

    def test_filtro_numerico_igual(self):
        from Sapiens_MultiAgente.tools.csv_search_tool import CSVSearchTool
        path = self._make_csv(self.CSV_BASICO)
        try:
            result = CSVSearchTool()._run(file_path=path, numeric_filters={"nota": "==7.0"})
            assert "Bob" in result
        finally:
            os.unlink(path)

    def test_filtro_numerico_diferente(self):
        from Sapiens_MultiAgente.tools.csv_search_tool import CSVSearchTool
        path = self._make_csv(self.CSV_BASICO)
        try:
            result = CSVSearchTool()._run(file_path=path, numeric_filters={"nota": "!=7.0"})
            assert "Alice" in result
            assert "Carla" in result
        finally:
            os.unlink(path)

    def test_filtro_numerico_sem_operador(self):
        from Sapiens_MultiAgente.tools.csv_search_tool import CSVSearchTool
        path = self._make_csv(self.CSV_BASICO)
        try:
            result = CSVSearchTool()._run(file_path=path, numeric_filters={"nota": "7.0"})
            assert "Bob" in result
        finally:
            os.unlink(path)

    def test_max_results_limita_saida(self):
        from Sapiens_MultiAgente.tools.csv_search_tool import CSVSearchTool
        linhas = ["nome,valor"] + [f"item{i},{i}" for i in range(100)]
        path = self._make_csv("\n".join(linhas))
        try:
            result = CSVSearchTool()._run(file_path=path, max_results=5)
            assert "RESULTADOS DA BUSCA" in result
            assert "primeiros 5" in result
        finally:
            os.unlink(path)

    def test_estatisticas_colunas_numericas(self):
        from Sapiens_MultiAgente.tools.csv_search_tool import CSVSearchTool
        path = self._make_csv(self.CSV_BASICO)
        try:
            result = CSVSearchTool()._run(file_path=path)
            assert "ESTATÍSTICAS" in result
        finally:
            os.unlink(path)

    def test_format_filters_sem_filtros(self):
        from Sapiens_MultiAgente.tools.csv_search_tool import CSVSearchTool
        tool = CSVSearchTool()
        result = tool._format_filters(None, None, None)
        assert "Nenhum filtro" in result

    def test_format_filters_com_todos(self):
        from Sapiens_MultiAgente.tools.csv_search_tool import CSVSearchTool
        tool = CSVSearchTool()
        result = tool._format_filters("texto", {"col": "val"}, {"num": ">5"})
        assert "texto" in result
        assert "col" in result
        assert "num" in result


# ---------------------------------------------------------------------------
# PDFSearchTool — testes via mocks do PdfReader (33% → ~70%)
# ---------------------------------------------------------------------------

class TestPDFSearchToolExtended:

    def test_arquivo_inexistente(self):
        from Sapiens_MultiAgente.tools.pdf_search_tool import PDFSearchTool
        result = PDFSearchTool()._run(file_path="/nao/existe.pdf")
        assert "ERRO" in result

    def test_parse_page_range_none(self):
        from Sapiens_MultiAgente.tools.pdf_search_tool import PDFSearchTool
        result = PDFSearchTool()._parse_page_range(None, 5)
        assert result == [0, 1, 2, 3, 4]

    def test_parse_page_range_all(self):
        from Sapiens_MultiAgente.tools.pdf_search_tool import PDFSearchTool
        result = PDFSearchTool()._parse_page_range("all", 5)
        assert result == [0, 1, 2, 3, 4]

    def test_parse_page_range_intervalo(self):
        from Sapiens_MultiAgente.tools.pdf_search_tool import PDFSearchTool
        result = PDFSearchTool()._parse_page_range("2-4", 5)
        assert result == [1, 2, 3]

    def test_parse_page_range_pagina_unica(self):
        from Sapiens_MultiAgente.tools.pdf_search_tool import PDFSearchTool
        result = PDFSearchTool()._parse_page_range("2", 5)
        assert result == [1]

    def test_parse_page_range_invalido(self):
        from Sapiens_MultiAgente.tools.pdf_search_tool import PDFSearchTool
        result = PDFSearchTool()._parse_page_range("abc", 3)
        assert result == [0, 1, 2]

    def test_extract_full_text_com_conteudo(self):
        from Sapiens_MultiAgente.tools.pdf_search_tool import PDFSearchTool
        tool = PDFSearchTool()
        mock_page = MagicMock()
        mock_page.extract_text.return_value = "Conteúdo acadêmico da página"
        mock_reader = MagicMock()
        mock_reader.pages = [mock_page, mock_page]
        result = tool._extract_full_text(mock_reader, [0, 1])
        assert "TEXTO EXTRAÍDO DO PDF" in result
        assert "Conteúdo acadêmico" in result

    def test_extract_full_text_sem_texto(self):
        from Sapiens_MultiAgente.tools.pdf_search_tool import PDFSearchTool
        tool = PDFSearchTool()
        mock_page = MagicMock()
        mock_page.extract_text.return_value = "   "
        mock_reader = MagicMock()
        mock_reader.pages = [mock_page]
        result = tool._extract_full_text(mock_reader, [0])
        assert "NENHUM TEXTO EXTRAÍDO" in result

    def test_search_text_encontrado(self):
        from Sapiens_MultiAgente.tools.pdf_search_tool import PDFSearchTool
        tool = PDFSearchTool()
        mock_page = MagicMock()
        mock_page.extract_text.return_value = "O desempenho acadêmico é excelente.\nOutra linha de dados."
        mock_reader = MagicMock()
        mock_reader.pages = [mock_page]
        result = tool._search_text(mock_reader, "acadêmico", [0])
        assert "RESULTADOS DA BUSCA" in result
        assert "PÁGINA 1" in result

    def test_search_text_nao_encontrado(self):
        from Sapiens_MultiAgente.tools.pdf_search_tool import PDFSearchTool
        tool = PDFSearchTool()
        mock_page = MagicMock()
        mock_page.extract_text.return_value = "Texto sem a palavra procurada."
        mock_reader = MagicMock()
        mock_reader.pages = [mock_page]
        result = tool._search_text(mock_reader, "XYZXYZ_inexistente", [0])
        assert "TERMO NÃO ENCONTRADO" in result

    def test_get_pdf_info(self):
        from Sapiens_MultiAgente.tools.pdf_search_tool import PDFSearchTool
        tool = PDFSearchTool()
        mock_page = MagicMock()
        mock_page.extract_text.return_value = "Preview do conteúdo da primeira página."
        mock_reader = MagicMock()
        mock_reader.pages = [mock_page]
        mock_reader.metadata = None
        result = tool._get_pdf_info(mock_reader, "/tmp/teste.pdf")
        assert "INFORMAÇÕES DO PDF" in result
        assert "teste.pdf" in result

    def test_run_extract_text_via_mock(self, tmp_path):
        from Sapiens_MultiAgente.tools.pdf_search_tool import PDFSearchTool
        pdf_path = str(tmp_path / "test.pdf")
        with open(pdf_path, 'wb') as f:
            f.write(b'%PDF-1.4')

        mock_page = MagicMock()
        mock_page.extract_text.return_value = "Texto extraído do PDF."
        mock_reader = MagicMock()
        mock_reader.pages = [mock_page]
        mock_reader.metadata = None

        with patch('Sapiens_MultiAgente.tools.pdf_search_tool.PdfReader', return_value=mock_reader):
            result = PDFSearchTool()._run(file_path=pdf_path, extract_text=True)
        assert "TEXTO EXTRAÍDO" in result

    def test_run_search_via_mock(self, tmp_path):
        from Sapiens_MultiAgente.tools.pdf_search_tool import PDFSearchTool
        pdf_path = str(tmp_path / "test2.pdf")
        with open(pdf_path, 'wb') as f:
            f.write(b'%PDF-1.4')

        mock_page = MagicMock()
        mock_page.extract_text.return_value = "Desempenho dos estudantes de graduação."
        mock_reader = MagicMock()
        mock_reader.pages = [mock_page]
        mock_reader.metadata = None

        with patch('Sapiens_MultiAgente.tools.pdf_search_tool.PdfReader', return_value=mock_reader):
            result = PDFSearchTool()._run(file_path=pdf_path, search_query="graduação")
        assert "RESULTADOS" in result or "PÁGINA" in result

    def test_run_info_via_mock(self, tmp_path):
        from Sapiens_MultiAgente.tools.pdf_search_tool import PDFSearchTool
        pdf_path = str(tmp_path / "test3.pdf")
        with open(pdf_path, 'wb') as f:
            f.write(b'%PDF-1.4')

        mock_page = MagicMock()
        mock_page.extract_text.return_value = "Preview do conteúdo."
        mock_reader = MagicMock()
        mock_reader.pages = [mock_page]
        mock_reader.metadata = None

        with patch('Sapiens_MultiAgente.tools.pdf_search_tool.PdfReader', return_value=mock_reader):
            result = PDFSearchTool()._run(file_path=pdf_path)
        assert "INFORMAÇÕES" in result


# ---------------------------------------------------------------------------
# DOCXSearchTool (25% → ~75%)
# ---------------------------------------------------------------------------

class TestDOCXSearchToolExtended:

    def _make_docx(self, paragraphs, tmp_path, with_table=False):
        try:
            from docx import Document
        except ImportError:
            pytest.skip("python-docx não disponível")

        doc = Document()
        for para in paragraphs:
            doc.add_paragraph(para)
        if with_table:
            table = doc.add_table(rows=2, cols=2)
            table.rows[0].cells[0].text = "Cabeçalho A"
            table.rows[0].cells[1].text = "Cabeçalho B"
            table.rows[1].cells[0].text = "Dado importante"
            table.rows[1].cells[1].text = "Dado 2"
        path = str(tmp_path / "test.docx")
        doc.save(path)
        return path

    def test_arquivo_inexistente(self):
        from Sapiens_MultiAgente.tools.docx_search_tool import DOCXSearchTool
        result = DOCXSearchTool()._run(file_path="/nao/existe.docx")
        assert "ERRO" in result

    def test_extract_text(self, tmp_path):
        from Sapiens_MultiAgente.tools.docx_search_tool import DOCXSearchTool
        path = self._make_docx(
            ["Primeiro parágrafo.", "Segundo parágrafo.", "Terceiro parágrafo."],
            tmp_path,
        )
        result = DOCXSearchTool()._run(file_path=path, extract_text=True)
        assert "TEXTO EXTRAÍDO" in result
        assert "Primeiro parágrafo" in result

    def test_extract_text_com_paragraph_limit(self, tmp_path):
        from Sapiens_MultiAgente.tools.docx_search_tool import DOCXSearchTool
        path = self._make_docx([f"Parágrafo {i}." for i in range(10)], tmp_path)
        result = DOCXSearchTool()._run(file_path=path, extract_text=True, paragraph_limit=3)
        assert "TEXTO EXTRAÍDO" in result

    def test_extract_text_doc_vazio(self, tmp_path):
        try:
            from docx import Document
        except ImportError:
            pytest.skip("python-docx não disponível")
        doc = Document()
        path = str(tmp_path / "vazio.docx")
        doc.save(path)
        from Sapiens_MultiAgente.tools.docx_search_tool import DOCXSearchTool
        result = DOCXSearchTool()._run(file_path=path, extract_text=True)
        assert "NENHUM TEXTO" in result

    def test_search_text_encontrado(self, tmp_path):
        from Sapiens_MultiAgente.tools.docx_search_tool import DOCXSearchTool
        path = self._make_docx(
            ["O desempenho acadêmico é excelente.", "Outro parágrafo sem a palavra."],
            tmp_path,
        )
        result = DOCXSearchTool()._run(file_path=path, search_query="acadêmico")
        assert "RESULTADOS DA BUSCA" in result

    def test_search_text_nao_encontrado(self, tmp_path):
        from Sapiens_MultiAgente.tools.docx_search_tool import DOCXSearchTool
        path = self._make_docx(["Texto sem a palavra procurada."], tmp_path)
        result = DOCXSearchTool()._run(file_path=path, search_query="XYZXYZ_inexistente")
        assert "TERMO NÃO ENCONTRADO" in result

    def test_search_em_tabela(self, tmp_path):
        from Sapiens_MultiAgente.tools.docx_search_tool import DOCXSearchTool
        path = self._make_docx(["Parágrafo sem a palavra."], tmp_path, with_table=True)
        result = DOCXSearchTool()._run(file_path=path, search_query="importante")
        assert "RESULTADOS" in result

    def test_get_info_basico(self, tmp_path):
        from Sapiens_MultiAgente.tools.docx_search_tool import DOCXSearchTool
        path = self._make_docx(["Parágrafo de teste."], tmp_path)
        result = DOCXSearchTool()._run(file_path=path)
        assert "INFORMAÇÕES DO DOCUMENTO DOCX" in result

    def test_get_info_com_tabela(self, tmp_path):
        from Sapiens_MultiAgente.tools.docx_search_tool import DOCXSearchTool
        path = self._make_docx(["Parágrafo."], tmp_path, with_table=True)
        result = DOCXSearchTool()._run(file_path=path)
        assert "Tabela" in result
