import os
from typing import Type, Optional, List
from pydantic import BaseModel, Field
from crewai.tools import BaseTool
import re
try:
    from docx import Document
except ImportError:
    # Fallback se python-docx não estiver disponível
    Document = None


class DOCXSearchToolInput(BaseModel):
    """Input schema for DOCXSearchTool."""
    file_path: str = Field(..., description="Path to the DOCX file to search")
    search_query: Optional[str] = Field(default=None, description="Text to search for (case-insensitive)")
    extract_text: bool = Field(default=False, description="Extract all text from the DOCX")
    extract_images: bool = Field(default=False, description="Extract images embedded in the DOCX and save to a directory")
    output_dir: Optional[str] = Field(default=None, description="Directory to save extracted images (defaults to temp dir)")
    paragraph_limit: Optional[int] = Field(default=None, description="Limit number of paragraphs to extract")


class DOCXSearchTool(BaseTool):
    name: str = "DOCXSearchTool"
    description: str = (
        "Ferramenta para buscar e extrair texto de arquivos DOCX. Permite pesquisa "
        "de termos específicos, extração de texto completo ou por parágrafos, "
        "e análise de conteúdo em documentos acadêmicos Word."
    )
    args_schema: Type[BaseModel] = DOCXSearchToolInput

    def _run(self, file_path: str, search_query: Optional[str] = None,
             extract_text: bool = False, extract_images: bool = False,
             output_dir: Optional[str] = None, paragraph_limit: Optional[int] = None) -> str:
        """
        Busca ou extrai texto de um arquivo DOCX.

        Args:
            file_path: Caminho para o arquivo DOCX
            search_query: Termo a buscar (opcional)
            extract_text: Se deve extrair todo o texto
            paragraph_limit: Limite de parágrafos para extrair

        Returns:
            Resultados da busca ou extração
        """
        if not os.path.exists(file_path):
            return f"❌ ERRO: Arquivo DOCX não encontrado: {file_path}"

        if Document is None:
            return "❌ ERRO: Biblioteca python-docx não está instalada. Instale com: pip install python-docx"

        try:
            # Abrir documento DOCX
            doc = Document(file_path)

            # Extrair imagens se solicitado
            if extract_images:
                return self._extract_images(file_path, output_dir)

            # Extrair texto se solicitado
            if extract_text:
                return self._extract_full_text(doc, paragraph_limit)

            # Buscar termo específico
            if search_query:
                return self._search_text(doc, search_query, paragraph_limit)

            # Se nada especificado, retornar informações básicas
            return self._get_docx_info(doc, file_path)

        except Exception as e:
            return f"❌ ERRO ao processar DOCX: {str(e)}"

    def _extract_full_text(self, doc: Document, paragraph_limit: Optional[int] = None) -> str:
        """Extrai todo o texto do documento DOCX."""
        paragraphs = []

        for i, paragraph in enumerate(doc.paragraphs):
            if paragraph_limit and i >= paragraph_limit:
                break

            text = paragraph.text.strip()
            if text:  # Só adicionar parágrafos não vazios
                paragraphs.append(f"§{i+1}: {text}")

        if not paragraphs:
            return "❌ NENHUM TEXTO EXTRAÍDO do DOCX. O documento pode estar vazio."

        result = f"📄 TEXTO EXTRAÍDO DO DOCUMENTO DOCX\n\n"
        result += f"📊 Total de parágrafos processados: {len(paragraphs)}\n\n"
        result += "\n\n".join(paragraphs)

        # Adicionar informações sobre tabelas se existirem
        if doc.tables:
            result += f"\n\n📋 Tabelas encontradas: {len(doc.tables)}"
            for i, table in enumerate(doc.tables):
                try:
                    rows, cols = len(table.rows), len(table.columns)
                    result += f"\n• Tabela {i+1}: {rows} linhas × {cols} colunas"
                except:
                    result += f"\n• Tabela {i+1}: erro ao processar"

        return result

    def _search_text(self, doc: Document, query: str, paragraph_limit: Optional[int] = None) -> str:
        """Busca por termo específico no documento DOCX."""
        results = []
        query_lower = query.lower()

        for i, paragraph in enumerate(doc.paragraphs):
            if paragraph_limit and i >= paragraph_limit:
                break

            text = paragraph.text
            if text and query_lower in text.lower():
                # Encontrar contexto ao redor da ocorrência
                # Dividir em sentenças para melhor contexto
                sentences = re.split(r'(?<=[.!?])\s+', text)
                matching_sentences = []

                for sentence in sentences:
                    if query_lower in sentence.lower():
                        matching_sentences.append(sentence.strip())

                if matching_sentences:
                    results.append(f"📍 PARÁGRAFO {i+1}:\n" + '\n'.join(f"• {sent}" for sent in matching_sentences))

        # Também buscar em tabelas
        table_results = self._search_in_tables(doc, query)
        if table_results:
            results.extend(table_results)

        if not results:
            return f"🔍 TERMO NÃO ENCONTRADO: '{query}' não foi encontrado no documento DOCX."

        total_occurrences = len(results)
        return f"🔍 RESULTADOS DA BUSCA POR '{query}' EM DOCX\n\nEncontrado em {total_occurrences} local(is):\n\n" + "\n\n".join(results)

    def _search_in_tables(self, doc: Document, query: str) -> List[str]:
        """Busca por termo específico nas tabelas do documento."""
        results = []
        query_lower = query.lower()

        for table_idx, table in enumerate(doc.tables):
            try:
                for row_idx, row in enumerate(table.rows):
                    for cell_idx, cell in enumerate(row.cells):
                        text = cell.text
                        if text and query_lower in text.lower():
                            results.append(f"📍 TABELA {table_idx+1}, LINHA {row_idx+1}, CÉLULA {cell_idx+1}:\n• {text.strip()}")
            except:
                continue

        return results

    def _get_docx_info(self, doc: Document, file_path: str) -> str:
        """Retorna informações básicas do documento DOCX."""
        # Contar parágrafos não vazios
        non_empty_paragraphs = sum(1 for p in doc.paragraphs if p.text.strip())

        # Informações sobre tabelas
        num_tables = len(doc.tables)
        table_info = ""
        if num_tables > 0:
            table_info = f"\n📋 Tabelas: {num_tables}"
            for i, table in enumerate(doc.tables):
                try:
                    rows, cols = len(table.rows), len(table.columns)
                    table_info += f"\n• Tabela {i+1}: {rows}×{cols}"
                except:
                    table_info += f"\n• Tabela {i+1}: erro ao processar"

        # Tentar obter propriedades do documento
        core_props = ""
        try:
            if hasattr(doc, 'core_properties'):
                props = doc.core_properties
                if props.title:
                    core_props += f"\n• Título: {props.title}"
                if props.author:
                    core_props += f"\n• Autor: {props.author}"
                if props.subject:
                    core_props += f"\n• Assunto: {props.subject}"
                if props.created:
                    core_props += f"\n• Criado em: {props.created}"
                if props.modified:
                    core_props += f"\n• Modificado em: {props.modified}"
        except:
            pass

        info = f"""
📝 INFORMAÇÕES DO DOCUMENTO DOCX

📁 Arquivo: {os.path.basename(file_path)}
📄 Parágrafos: {non_empty_paragraphs}{table_info}

📋 PROPRIEDADES DO DOCUMENTO:{core_props}
"""

        # Preview do início do documento
        try:
            preview_paragraphs = []
            for paragraph in doc.paragraphs[:3]:  # Primeiros 3 parágrafos
                text = paragraph.text.strip()
                if text:
                    preview_paragraphs.append(text[:200] + "..." if len(text) > 200 else text)

            if preview_paragraphs:
                info += f"\n👀 PREVIEW DO DOCUMENTO:\n" + "\n\n".join(f"§{i+1}: {text}" for i, text in enumerate(preview_paragraphs))
        except:
            info += "\n👀 Preview não disponível."

        return info

    def _extract_images(self, file_path: str, output_dir: Optional[str]) -> str:
        """Extrai imagens embutidas no DOCX e salva em disco."""
        import zipfile
        import tempfile

        dest = output_dir or tempfile.mkdtemp(prefix='sapiens_docx_imgs_')
        os.makedirs(dest, exist_ok=True)

        # DOCX é um ZIP — imagens ficam em word/media/
        saved = []
        errors = []

        try:
            with zipfile.ZipFile(file_path, 'r') as zf:
                media_files = [
                    name for name in zf.namelist()
                    if name.startswith('word/media/')
                    and not name.endswith('/')
                ]

                if not media_files:
                    return (
                        f"🖼️ EXTRAÇÃO DE IMAGENS DO DOCX\n"
                        f"ℹ️ Nenhuma imagem encontrada em '{os.path.basename(file_path)}'."
                    )

                for media_path in media_files:
                    fname = os.path.basename(media_path)
                    fpath = os.path.join(dest, fname)
                    try:
                        with zf.open(media_path) as src, open(fpath, 'wb') as dst:
                            dst.write(src.read())
                        size = os.path.getsize(fpath)
                        saved.append(f"• {fname} ({size:,} bytes)")
                    except Exception as e:
                        errors.append(f"  {fname}: {e}")

        except zipfile.BadZipFile:
            return f"❌ ERRO: '{os.path.basename(file_path)}' não é um arquivo DOCX válido."
        except Exception as e:
            return f"❌ ERRO ao extrair imagens: {e}"

        report = f"🖼️ EXTRAÇÃO DE IMAGENS DO DOCX\n"
        report += f"📁 Destino: {dest}\n\n"

        if saved:
            report += f"✅ {len(saved)} imagem(ns) extraída(s):\n" + "\n".join(saved) + "\n"
        else:
            report += "ℹ️ Nenhuma imagem foi extraída.\n"

        if errors:
            report += f"\n⚠️ {len(errors)} erro(s):\n" + "\n".join(errors) + "\n"

        return report